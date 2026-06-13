"""Prepare dashboard Ask? chat payloads with optional file attachments."""
from __future__ import annotations

import base64
import binascii
import io
import re
from typing import Any

MAX_ATTACHMENT_BYTES = 8 * 1024 * 1024
MAX_ATTACHMENTS = 6
MAX_IMAGES = 4
MAX_PDF_PAGES = 40

_IMAGE_MIME = {
    b'\xff\xd8\xff': 'image/jpeg',
    b'\x89PNG\r\n\x1a\n': 'image/png',
    b'GIF87a': 'image/gif',
    b'GIF89a': 'image/gif',
    b'RIFF': 'image/webp',  # WEBP after RIFF....WEBP — validated loosely
}


class AttachmentError(ValueError):
    """Invalid or unsupported attachment input."""


def _decode_b64(content: str) -> bytes:
    if not content or not isinstance(content, str):
        raise AttachmentError('Attachment content must be base64 text')
    try:
        raw = base64.b64decode(content, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise AttachmentError('Invalid base64 attachment data') from exc
    if len(raw) > MAX_ATTACHMENT_BYTES:
        raise AttachmentError(
            f'Attachment exceeds {MAX_ATTACHMENT_BYTES // (1024 * 1024)} MB limit'
        )
    if not raw:
        raise AttachmentError('Attachment is empty')
    return raw


def _guess_image(raw: bytes) -> None:
    head = raw[:12]
    if head.startswith(b'\xff\xd8\xff'):
        return
    if head.startswith(b'\x89PNG\r\n\x1a\n'):
        return
    if head[:6] in (b'GIF87a', b'GIF89a'):
        return
    if head.startswith(b'RIFF') and b'WEBP' in head:
        return
    raise AttachmentError('Unsupported image format (use JPEG, PNG, GIF, or WebP)')


def _extract_pdf_text(raw: bytes, name: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise AttachmentError('PDF support is not installed on the server (pypdf)') from exc
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise AttachmentError(f'Could not read PDF {name!r}') from exc
    pages = reader.pages[:MAX_PDF_PAGES]
    chunks: list[str] = []
    for idx, page in enumerate(pages, start=1):
        try:
            text = (page.extract_text() or '').strip()
        except Exception:
            text = ''
        if text:
            chunks.append(f'--- Page {idx} ---\n{text}')
    if not chunks:
        raise AttachmentError(
            f'No extractable text in {name!r}. Try a vision model with page screenshots as images.'
        )
    suffix = ''
    if len(reader.pages) > MAX_PDF_PAGES:
        suffix = f'\n[Truncated after {MAX_PDF_PAGES} pages]'
    return f'[PDF: {name}]\n' + '\n\n'.join(chunks) + suffix


def _extract_docx_text(raw: bytes, name: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise AttachmentError('Word support is not installed on the server (python-docx)') from exc
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        raise AttachmentError(f'Could not read Word document {name!r}') from exc
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    text = '\n'.join(parts).strip()
    if not text:
        raise AttachmentError(f'No extractable text in {name!r}')
    return f'[Document: {name}]\n{text}'


def _extract_legacy_doc_text(raw: bytes, name: str) -> str:
    """Best-effort text recovery for binary .doc (not .docx)."""
    if raw[:2] == b'PK':
        return _extract_docx_text(raw, name)
    # Pull printable runs from common UTF-16LE / ASCII pockets in legacy .doc binaries.
    utf16 = raw.decode('utf-16-le', errors='ignore')
    ascii_runs = re.findall(r'[\x20-\x7e\n\r\t]{4,}', raw.decode('latin-1', errors='ignore'))
    utf16_runs = re.findall(r'[^\x00-\x08\x0b\x0c\x0e-\x1f]{4,}', utf16)
    merged = '\n'.join(dict.fromkeys([*utf16_runs[:80], *ascii_runs[:80]])).strip()
    if len(merged) < 20:
        raise AttachmentError(
            f'Could not extract text from legacy {name!r}. Save as .docx and try again.'
        )
    return f'[Document: {name}]\n{merged[:120_000]}'


def _format_code_block(content: str, language: str | None, name: str | None) -> str:
    lang = (language or 'text').strip() or 'text'
    label = name or f'code.{lang}'
    body = content.rstrip()
    return f'[Code snippet: {label}]\n```{lang}\n{body}\n```'


def prepare_chat_from_attachments(
    prompt: str,
    attachments: list[dict[str, Any]] | None,
    *,
    model_has_vision: bool | None = None,
) -> dict[str, Any]:
    """
    Merge user prompt with attachments for Ollama ``/api/generate``.

    Returns ``{"prompt": str, "images": list[str] | None}``.
    """
    prompt = (prompt or '').strip()
    items = attachments if isinstance(attachments, list) else []
    if len(items) > MAX_ATTACHMENTS:
        raise AttachmentError(f'Maximum {MAX_ATTACHMENTS} attachments allowed')

    doc_blocks: list[str] = []
    images: list[str] = []

    for item in items:
        if not isinstance(item, dict):
            raise AttachmentError('Each attachment must be an object')
        kind = str(item.get('type') or '').strip().lower()
        name = str(item.get('name') or 'attachment').strip() or 'attachment'

        if kind == 'code':
            text = str(item.get('content') or item.get('text') or '').strip()
            if not text:
                raise AttachmentError('Code snippet is empty')
            doc_blocks.append(_format_code_block(text, item.get('language'), name))
            continue

        raw: bytes | None = None
        b64 = item.get('content')
        if isinstance(b64, str) and b64.strip():
            raw = _decode_b64(b64)
        elif kind != 'code':
            raise AttachmentError(f'Missing content for attachment {name!r}')

        if kind == 'image':
            if model_has_vision is False:
                raise AttachmentError(
                    f'Model does not support images. Remove {name!r} or use a vision-capable model.'
                )
            _guess_image(raw)
            if len(images) >= MAX_IMAGES:
                raise AttachmentError(f'Maximum {MAX_IMAGES} images per message')
            images.append(base64.b64encode(raw).decode('ascii'))
            continue

        if kind == 'pdf':
            doc_blocks.append(_extract_pdf_text(raw, name))
            continue

        if kind in ('doc', 'docx', 'word'):
            lower = name.lower()
            if lower.endswith('.docx') or raw[:2] == b'PK':
                doc_blocks.append(_extract_docx_text(raw, name))
            elif lower.endswith('.doc'):
                doc_blocks.append(_extract_legacy_doc_text(raw, name))
            else:
                doc_blocks.append(_extract_docx_text(raw, name))
            continue

        raise AttachmentError(f'Unsupported attachment type: {kind!r}')

    if doc_blocks:
        joined = '\n\n'.join(doc_blocks)
        prompt = f'{joined}\n\n{prompt}' if prompt else joined

    if not prompt and not images:
        raise AttachmentError('Enter a question or add an attachment')

    if not prompt and images:
        prompt = 'Describe the attached image(s) in detail.'

    return {
        'prompt': prompt,
        'images': images or None,
    }
